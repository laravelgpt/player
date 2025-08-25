import os
import threading
import fitz  # PyMuPDF
import pandas as pd
from PIL import Image
from docx import Document
from pdf2docx import Converter
from tkinter import filedialog, messagebox, Tk, ttk
from tqdm import tqdm

class FileConverter:
    def __init__(self, root):
        self.root = root
        self.selected_file = None
        self.format_var = None
        self.progress = ttk.Progressbar(root, orient='horizontal', length=300, mode='determinate')
        self.progress.pack()
    
    def select_file(self):
        self.selected_file = filedialog.askopenfilename()
    
    def convert_file(self):
        if not self.selected_file:
            messagebox.showerror("Error", "No file selected!")
            return

        target_format = self.format_var.get()
        output_file = filedialog.asksaveasfilename(defaultextension=f".{target_format}")

        if output_file:
            try:
                threading.Thread(target=self.process_conversion, args=(target_format, output_file)).start()
            except Exception as e:
                messagebox.showerror("Error", f"Conversion failed: {e}")

    def process_conversion(self, target_format, output_file):
        self.progress['value'] = 0
        
        ext = os.path.splitext(self.selected_file)[1].lower()
        
        conversion_map = {
            ".pdf": self.convert_pdf,
            ".docx": self.convert_docx,
            ".txt": self.convert_txt,
            ".jpg": self.convert_image,
            ".png": self.convert_image,
            ".gif": self.convert_image,
            ".mp4": self.convert_video,
            ".avi": self.convert_video,
            ".mov": self.convert_video,
            ".mkv": self.convert_video,
            ".php": self.convert_code,
            ".py": self.convert_code,
            ".c": self.convert_code,
            ".cs": self.convert_code,
            ".vb": self.convert_code
        }

        if ext in conversion_map:
            conversion_map[ext](target_format, output_file)
        else:
            messagebox.showerror("Error", "Unsupported file format!")
        
        self.progress['value'] = 100
        messagebox.showinfo("Success", f"{target_format.upper()} saved successfully!")
    
    def convert_pdf(self, target_format, output_file):
        doc = fitz.open(self.selected_file)
        
        if target_format == "docx":
            cv = Converter(self.selected_file)
            cv.convert(output_file, start=0, end=None)
            cv.close()
        elif target_format in ["xls", "xlsx"]:
            text = "".join([page.get_text() for page in doc])
            df = pd.DataFrame([text.split("\n")])
            df.to_excel(output_file, index=False)
        elif target_format in ["jpg", "png"]:
            page = doc.load_page(0)
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            img.save(output_file)
        else:
            raise ValueError("Unsupported target format for PDF!")
    
    def convert_docx(self, target_format, output_file):
        doc = Document(self.selected_file)
        
        if target_format == "pdf":
            doc.save(output_file)
        elif target_format == "txt":
            with open(output_file, "w", encoding="utf-8") as file:
                file.write("\n".join([para.text for para in doc.paragraphs]))
        else:
            raise ValueError("Unsupported target format for DOCX!")
    
    def convert_txt(self, target_format, output_file):
        with open(self.selected_file, "r", encoding="utf-8") as file:
            content = file.readlines()
        
        if target_format == "docx":
            doc = Document()
            for line in content:
                doc.add_paragraph(line.strip())
            doc.save(output_file)
        elif target_format in ["xls", "xlsx", "csv"]:
            df = pd.DataFrame([content])
            getattr(df, f"to_{target_format}")(output_file, index=False)
        elif target_format in ["jpg", "png", "gif"]:
            img = Image.new("RGB", (800, 600), color="white")
            img.save(output_file)
        else:
            raise ValueError("Unsupported target format for TXT!")
    
    def convert_image(self, target_format, output_file):
        img = Image.open(self.selected_file)
        img.save(output_file)
    
    def convert_video(self, target_format, output_file):
        messagebox.showerror("Error", "Video conversion not implemented!")
    
    def convert_code(self, target_format, output_file):
        with open(self.selected_file, "r", encoding="utf-8") as file:
            code = file.read()
        
        if target_format in ["php", "py", "c", "cs", "vb"]:
            with open(output_file, "w", encoding="utf-8") as out_file:
                out_file.write(code)
        else:
            raise ValueError("Unsupported target format for code!")

if __name__ == "__main__":
    root = Tk()
    app = FileConverter(root)
    root.mainloop()
